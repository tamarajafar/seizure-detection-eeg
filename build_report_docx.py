"""
Build report.docx from report.md, embedding figures at the correct locations.
Run from repo root: python3 build_report_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT = Path("report.md")
FIGURES_DIR = Path("results/figures")
OUT = Path("report.docx")

FIGURE_FILES = {
    "Figure 1": FIGURES_DIR / "fig1_auroc_comparison.png",
    "Figure 2": FIGURES_DIR / "fig2_metrics_comparison.png",
    "Figure 3": FIGURES_DIR / "fig3_per_subject_auroc.png",
    "Figure 4": FIGURES_DIR / "fig4_roc_curves.png",
    "Figure 5": FIGURES_DIR / "fig5_confusion_matrices.png",
}


def set_font(run, bold=False, italic=False, size=None, color=None):
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_figure(doc, fig_key, caption_text):
    """Insert figure image then caption paragraph."""
    fig_path = FIGURE_FILES.get(fig_key)
    if fig_path and fig_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fig_path), width=Inches(5.5))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bold_run = cap.add_run(f"{fig_key}. ")
    bold_run.bold = True
    bold_run.font.size = Pt(9)
    rest = cap.add_run(caption_text.strip())
    rest.font.size = Pt(9)
    rest.italic = True
    doc.add_paragraph()  # spacing


def parse_inline(paragraph, text):
    """Add runs to paragraph handling **bold**, *italic*, and `code`."""
    # Split on bold, italic, backtick patterns
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*'):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Courier New'
        else:
            paragraph.add_run(part)


def add_table_from_md(doc, rows):
    """Build a Word table from list of pipe-delimited row strings."""
    # Filter out separator rows (---|---|---)
    data_rows = [r for r in rows if not re.match(r'^\s*\|[-:\s|]+\|\s*$', r)]
    if not data_rows:
        return
    # Parse cells
    table_data = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        table_data.append(cells)

    n_cols = max(len(r) for r in table_data)
    table = doc.add_table(rows=len(table_data), cols=n_cols)
    table.style = 'Table Grid'

    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                break
            cell = table.cell(i, j)
            # Strip markdown bold
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(clean)
            run.font.size = Pt(9)
            if i == 0:  # header row
                run.bold = True
                cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D9D9D9')
                cell._tc.tcPr.append(shd)
    doc.add_paragraph()


def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = REPORT.read_text().splitlines()

    # Figure caption lines accumulate separately
    figure_caption_pattern = re.compile(r'^\*\*(Figure \d+)\.\*\*\s*(.*)')

    i = 0
    table_buffer = []
    in_table = False
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # --- Code block ---
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                for cl in code_lines:
                    run = p.add_run(cl + '\n')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
                doc.add_paragraph()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Table row ---
        if line.strip().startswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        else:
            if table_buffer:
                add_table_from_md(doc, table_buffer)
                table_buffer = []

        # --- Blank line ---
        if not line.strip():
            i += 1
            continue

        # --- Horizontal rule ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Headings ---
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # --- Bold metadata lines (Authors, course, etc.) ---
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.strip('*'))
            r.bold = True
            i += 1
            continue

        # --- Figure caption lines ---
        fig_match = figure_caption_pattern.match(line)
        if fig_match:
            fig_key = fig_match.group(1)
            caption_text = fig_match.group(2)
            add_figure(doc, fig_key, caption_text)
            i += 1
            continue

        # --- Bullet points ---
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, line[2:].strip())
            i += 1
            continue

        # --- Numbered list ---
        if re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, re.sub(r'^\d+\.\s', '', line))
            i += 1
            continue

        # --- Block quote / indented ---
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            parse_inline(p, line[2:])
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        parse_inline(p, line.strip())
        i += 1

    # Flush any remaining table
    if table_buffer:
        add_table_from_md(doc, table_buffer)

    doc.save(str(OUT))
    print(f"Saved {OUT}")


if __name__ == '__main__':
    build_docx()
